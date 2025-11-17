# core/document_generator.py
import os
from pathlib import Path
from django.conf import settings
from django.template.loader import render_to_string
from django.utils import timezone

from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

try:
    from weasyprint import HTML  # optional
    HAS_WEASY = True
except Exception:
    HAS_WEASY = False

from .models import Pagamento

DEFAULT_OUTPUT = getattr(settings, 'MEDIA_ROOT', Path.cwd() / 'media')
RECIBOS_SUBDIR = 'recibos'

class DocumentGenerator:
    def __init__(self, output_dir: str = None):
        base = output_dir or DEFAULT_OUTPUT
        self.output_dir = os.path.join(str(base), RECIBOS_SUBDIR)
        os.makedirs(self.output_dir, exist_ok=True)

    def _sanitize_filename(self, name: str) -> str:
        return "".join(c for c in name if c.isalnum() or c in (' ', '-', '_', '.')).rstrip()

    def gerar_recibo_pagamento(self, pagamento_id) -> str:
        pagamento = Pagamento.objects.select_related(
            'comanda', 'comanda__locacao', 'comanda__locacao__locatario', 'comanda__locacao__imovel'
        ).get(pk=pagamento_id)

        numero = pagamento.numero_pagamento or str(pagamento_id)
        timestamp = timezone.now().strftime('%Y%m%d%H%M%S')
        filename = f"recibo_{numero}_{timestamp}.docx"
        filename = self._sanitize_filename(filename)
        path = os.path.join(self.output_dir, filename)

        doc = Document()
        section = doc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run("RECIBO DE PAGAMENTO\n")
        run.bold = True
        run.font.size = Pt(16)

        doc.add_paragraph()

        def add_kv(label, value):
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            run.font.size = Pt(11)
            run2 = p.add_run(str(value or "-"))
            run2.font.size = Pt(11)

        locatario = getattr(pagamento.comanda.locacao, 'locatario', None)
        imovel = getattr(pagamento.comanda.locacao, 'imovel', None)

        add_kv("Locatário", getattr(locatario, 'nome_razao_social', '-'))
        add_kv("CPF/CNPJ", getattr(locatario, 'cpf_cnpj', '-'))
        add_kv("Imóvel", f"{getattr(imovel, 'endereco', '-')}, {getattr(imovel, 'numero', '')}")
        add_kv("Comanda", pagamento.comanda.numero_comanda if pagamento.comanda else '-')
        add_kv("Recibo", numero)
        add_kv("Data do Pagamento", pagamento.data_pagamento.strftime('%d/%m/%Y') if pagamento.data_pagamento else '-')
        add_kv("Forma", pagamento.get_forma_pagamento_display() if hasattr(pagamento, 'get_forma_pagamento_display') else pagamento.forma_pagamento)
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        valor_str = f"R$ {pagamento.valor_pago:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        run = p.add_run(f"VALOR PAGO\n{valor_str}")
        run.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph()

        if pagamento.observacoes:
            add_kv("Observações", pagamento.observacoes)

        doc.add_paragraph()
        doc.add_paragraph("Atenciosamente,")
        doc.add_paragraph(getattr(settings, 'DEFAULT_FROM_EMAIL', 'HABITAT PRO'))

        doc.save(path)
        return os.path.basename(path)

    def gerar_recibo_pdf(self, pagamento_id) -> str:
        if not HAS_WEASY:
            raise RuntimeError("WeasyPrint não está disponível. Instale weasyprint para gerar PDF.")

        pagamento = Pagamento.objects.select_related(
            'comanda', 'comanda__locacao', 'comanda__locacao__locatario', 'comanda__locacao__imovel'
        ).get(pk=pagamento_id)
        context = {
            'pagamento': pagamento,
            'comanda': pagamento.comanda,
            'locacao': pagamento.comanda.locacao if pagamento.comanda else None,
            'locatario': pagamento.comanda.locacao.locatario if pagamento.comanda and pagamento.comanda.locacao else None,
            'imovel': pagamento.comanda.locacao.imovel if pagamento.comanda and pagamento.comanda.locacao else None,
        }
        html = render_to_string('admin/pagamento_recibo.html', context)
        timestamp = timezone.now().strftime('%Y%m%d%H%M%S')
        filename = f"recibo_{pagamento.numero_pagamento}_{timestamp}.pdf"
        filename = self._sanitize_filename(filename)
        path = os.path.join(self.output_dir, filename)
        HTML(string=html, base_url=settings.STATIC_ROOT or settings.STATIC_URL).write_pdf(path)
        return os.path.basename(path)
