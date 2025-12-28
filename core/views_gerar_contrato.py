"""
View para geração de contratos de locação
HABITAT PRO - Sistema de Gestão Imobiliária
"""
from django.shortcuts import get_object_or_404, redirect
from django.http import HttpResponse
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib import messages
from datetime import datetime
import io
import os
import tempfile
import subprocess

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import Locacao, TemplateContrato


def formatar_cpf_cnpj(valor):
    """Formata CPF ou CNPJ."""
    if not valor:
        return 'Não informado'
    
    numeros = ''.join(filter(str.isdigit, valor))
    
    if len(numeros) == 11:  # CPF
        return f"{numeros[:3]}.{numeros[3:6]}.{numeros[6:9]}-{numeros[9:]}"
    elif len(numeros) == 14:  # CNPJ
        return f"{numeros[:2]}.{numeros[2:5]}.{numeros[5:8]}/{numeros[8:12]}-{numeros[12:]}"
    
    return valor


def formatar_data(data):
    """Formata data para o padrão brasileiro."""
    if not data:
        return 'Não informado'
    
    if isinstance(data, str):
        return data
    
    return data.strftime('%d/%m/%Y')


def formatar_moeda(valor):
    """Formata valor monetário."""
    if not valor:
        return 'R$ 0,00'
    
    return f"R$ {float(valor):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')


def buscar_template_contrato(locacao):
    """Busca o template de contrato mais adequado."""
    locador = locacao.imovel.locador
    tipo_imovel = locacao.imovel.tipo_imovel
    
    # 1. Específico do locador + tipo
    template = TemplateContrato.objects.filter(
        locador=locador,
        tipo_imovel=tipo_imovel,
        is_active=True
    ).first()
    
    if template:
        return template
    
    # 2. Específico do locador
    template = TemplateContrato.objects.filter(
        locador=locador,
        tipo_imovel='',
        is_active=True
    ).first()
    
    if template:
        return template
    
    # 3. Específico do tipo de imóvel
    template = TemplateContrato.objects.filter(
        locador=None,
        tipo_imovel=tipo_imovel,
        is_active=True
    ).first()
    
    if template:
        return template
    
    # 4. Template padrão
    template = TemplateContrato.objects.filter(
        is_default=True,
        is_active=True
    ).first()
    
    return template


def preparar_contexto_contrato(locacao):
    """
    Prepara o contexto com todas as variáveis do contrato.
    ✅ CORRIGIDO DEV_20: Adiciona variáveis de garantia (tipo_garantia, caução, seguro, fiador_garantia)
    """
    locador = locacao.imovel.locador
    locatario = locacao.locatario
    imovel = locacao.imovel
    # ✅ CORREÇÃO: Pegar fiador da LOCAÇÃO (garantia), não do locatário
    fiador = locacao.fiador_garantia
    
    contexto = {
        # LOCADOR
        'locador_nome': locador.nome_razao_social or 'Não informado',
        'locador_endereco_completo': locador.endereco_completo or 'Não informado',
        'locador_cep': locador.cep or 'Não informado',
        'locador_representante': locador.representante or locador.nome_razao_social or 'Não informado',
        'Locador_Representante': locador.representante or locador.nome_razao_social or 'Não informado',  # Case alternativo
        
        # LOCATÁRIO
        'locatario_nome': locatario.nome_razao_social or 'Não informado',
        'locatario_rg': locatario.rg or 'Não informado',
        'locatario_cpf_cnpj': formatar_cpf_cnpj(locatario.cpf_cnpj),
        'locatario_data_nascimento': formatar_data(locatario.data_nascimento),
        'locatario_telefone': locatario.telefone or 'Não informado',
        'locatario_filiacao_pai': locatario.nome_pai or 'Não informado',
        'locatario_filiacao_mae': locatario.nome_mae or 'Não informado',
        'locatario_endereco_completo': locatario.endereco_completo or 'Não informado',
        'locatario_fiador': fiador.nome_completo if fiador else 'Não possui fiador',
        
        # FIADOR - Dados Essenciais
        'fiador_nome': fiador.nome_completo if fiador else '',
        'fiador_cpf': formatar_cpf_cnpj(fiador.cpf) if fiador and fiador.cpf else '',
        'fiador_rg': fiador.rg if fiador and fiador.rg else '',
        'fiador_data_nascimento': formatar_data(fiador.data_nascimento) if fiador and fiador.data_nascimento else '',
        'fiador_telefone': fiador.telefone if fiador and fiador.telefone else '',
        'fiador_email': fiador.email if fiador and fiador.email else '',
        'fiador_endereco_completo': fiador.endereco_completo if fiador and fiador.endereco_completo else '',
        'fiador_cep': fiador.cep if fiador and fiador.cep else '',
        # ✅ NOVO DEV_20: Campos adicionais do fiador
        'fiador_profissao': '',  # Campo não existe no modelo
        'fiador_endereco': fiador.endereco_completo if fiador and fiador.endereco_completo else '',
        
        # IMÓVEL - Endereço
        'imovel_endereco': imovel.endereco or 'Não informado',
        'imovel_numero': imovel.numero or 'S/N',
        'imovel_bairro': imovel.bairro or 'Não informado',
        'imovel_cidade': imovel.cidade or 'Não informado',
        'imovel_estado': imovel.estado or 'PR',
        'imovel_cep': imovel.cep or 'Não informado',
        
        # IMÓVEL - Utilidades
        'imovel_conta_agua_esgoto': imovel.conta_agua_esgoto or 'Não informado',
        'imovel_numero_hidrometro': imovel.numero_hidrometro or 'Não informado',
        'imovel_unidade_consumidora_energia': imovel.unidade_consumidora_energia or 'Não informado',
        'imovel_descricao': imovel.descricao or 'Conforme vistoria anexa',
        
        # LOCAÇÃO - Valores
        'valor_aluguel': formatar_moeda(locacao.valor_aluguel),
        'locacao_data_inicio': formatar_data(locacao.data_inicio),
        'numero_contrato': locacao.numero_contrato or 'Não gerado',
        
        # ========================================
        # ✅ NOVO DEV_20: GARANTIAS
        # ========================================
        'tipo_garantia': dict(locacao.TIPO_GARANTIA_CHOICES).get(locacao.tipo_garantia, 'Não informado'),
        
        # Caução
        'caucao_meses': str(locacao.caucao_quantidade_meses) if locacao.caucao_quantidade_meses else '',
        'caucao_valor': formatar_moeda(locacao.caucao_valor_total) if locacao.caucao_valor_total else '',
        
        # Seguro
        'seguro_apolice': locacao.seguro_apolice or '',
        'seguro_seguradora': locacao.seguro_seguradora or '',
        
        # Extras
        'data_hoje': datetime.now().strftime('%d/%m/%Y'),
        'cidade': 'PARANAGUÁ',
    }
    
    return contexto


def adicionar_cabecalho_rodape(doc, numero_contrato):
    """
    Adiciona cabeçalho e rodapé personalizados em todas as páginas.
    Cabeçalho: Número do contrato (direita)
    Rodapé: HABITAT PRO v1.0 (centralizado, colorido)
    """
    for section in doc.sections:
        # ===== CABEÇALHO =====
        header = section.header
        
        # Limpar cabeçalho existente
        for paragraph in header.paragraphs:
            paragraph.clear()
        
        # Adicionar número do contrato (direita)
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run = paragraph.add_run(f'Contrato Nº {numero_contrato}')
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 58, 138)  # Azul HABITAT PRO
        
        # ===== RODAPÉ =====
        footer = section.footer
        
        # Limpar rodapé existente
        for paragraph in footer.paragraphs:
            paragraph.clear()
        
        # Adicionar texto do rodapé (centralizado)
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Texto: "Documento gerado via "
        run1 = paragraph.add_run('Documento gerado via ')
        run1.font.size = Pt(9)
        run1.font.color.rgb = RGBColor(100, 100, 100)  # Cinza
        
        # Texto: "HABITAT " (azul)
        run2 = paragraph.add_run('HABITAT ')
        run2.font.size = Pt(9)
        run2.font.bold = True
        run2.font.color.rgb = RGBColor(30, 58, 138)  # Azul
        
        # Texto: "PRO" (laranja)
        run3 = paragraph.add_run('PRO')
        run3.font.size = Pt(9)
        run3.font.bold = True
        run3.font.color.rgb = RGBColor(255, 140, 0)  # Laranja
        
        # Texto: " v1.0 | Sistema de Gestão Imobiliária"
        run4 = paragraph.add_run(' v1.0 | Sistema de Gestão Imobiliária')
        run4.font.size = Pt(9)
        run4.font.color.rgb = RGBColor(100, 100, 100)  # Cinza


def gerar_docx_contrato(locacao):
    """Gera o DOCX do contrato e retorna BytesIO."""
    # Buscar template
    template_obj = buscar_template_contrato(locacao)
    
    if not template_obj:
        raise Exception('Nenhum template de contrato encontrado.')
    
    # Preparar contexto
    contexto = preparar_contexto_contrato(locacao)
    
    # Carregar template
    doc_template = DocxTemplate(template_obj.arquivo_template.path)
    
    # Preencher template
    doc_template.render(contexto)
    
    # Salvar em memória
    docx_io = io.BytesIO()
    doc_template.save(docx_io)
    docx_io.seek(0)
    
    # Adicionar cabeçalho e rodapé personalizados
    doc = Document(docx_io)
    adicionar_cabecalho_rodape(doc, locacao.numero_contrato)
    
    # Salvar novamente
    final_io = io.BytesIO()
    doc.save(final_io)
    final_io.seek(0)
    
    return final_io


def converter_docx_para_pdf(docx_bytes):
    """
    Converte DOCX para PDF usando LibreOffice.
    Retorna BytesIO com o PDF ou None se falhar.
    """
    # Criar diretório temporário
    with tempfile.TemporaryDirectory() as temp_dir:
        # Salvar DOCX temporariamente
        docx_path = os.path.join(temp_dir, 'temp_contrato.docx')
        with open(docx_path, 'wb') as f:
            f.write(docx_bytes.read())
        
        try:
            # Converter usando LibreOffice
            subprocess.run([
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', temp_dir,
                docx_path
            ], check=True, timeout=30, capture_output=True)
            
            # Ler PDF gerado
            pdf_path = os.path.join(temp_dir, 'temp_contrato.pdf')
            
            if os.path.exists(pdf_path):
                with open(pdf_path, 'rb') as f:
                    pdf_bytes = io.BytesIO(f.read())
                return pdf_bytes
            else:
                return None
                
        except subprocess.TimeoutExpired:
            raise Exception('Timeout ao converter para PDF')
        except subprocess.CalledProcessError as e:
            raise Exception(f'Erro ao converter para PDF: {e.stderr.decode()}')
        except FileNotFoundError:
            raise Exception('LibreOffice não encontrado. Instale: sudo apt install libreoffice-writer')


@staff_member_required
def gerar_contrato_docx(request, locacao_id):
    """Gera contrato em formato DOCX."""
    locacao = get_object_or_404(Locacao, pk=locacao_id)
    
    try:
        # Gerar DOCX
        docx_io = gerar_docx_contrato(locacao)
        
        # Preparar resposta
        nome_arquivo = f'Contrato_{locacao.numero_contrato}_{locacao.locatario.nome_razao_social}'
        nome_arquivo = nome_arquivo.replace(' ', '_').replace('/', '-').replace('\\', '-')
        nome_arquivo = f'{nome_arquivo}.docx'
        
        response = HttpResponse(
            docx_io.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{nome_arquivo}"'
        
        messages.success(request, f'✅ Contrato DOCX gerado com sucesso!')
        
        return response
        
    except Exception as e:
        messages.error(request, f'❌ Erro ao gerar contrato: {str(e)}')
        return redirect('admin:core_locacao_change', locacao_id)


@staff_member_required
def gerar_contrato_pdf(request, locacao_id):
    """Gera contrato em formato PDF."""
    locacao = get_object_or_404(Locacao, pk=locacao_id)
    
    try:
        # Gerar DOCX
        docx_io = gerar_docx_contrato(locacao)
        
        # Converter para PDF
        pdf_io = converter_docx_para_pdf(docx_io)
        
        if not pdf_io:
            raise Exception('Falha ao gerar PDF')
        
        # Preparar resposta
        nome_arquivo = f'Contrato_{locacao.numero_contrato}_{locacao.locatario.nome_razao_social}'
        nome_arquivo = nome_arquivo.replace(' ', '_').replace('/', '-').replace('\\', '-')
        nome_arquivo = f'{nome_arquivo}.pdf'
        
        response = HttpResponse(
            pdf_io.read(),
            content_type='application/pdf'
        )
        response['Content-Disposition'] = f'attachment; filename="{nome_arquivo}"'
        
        messages.success(request, f'✅ Contrato PDF gerado com sucesso!')
        
        return response
        
    except Exception as e:
        messages.error(request, f'❌ Erro ao gerar PDF: {str(e)}')
        return redirect('admin:core_locacao_change', pk)


# ════════════════════════════════════════════════════════════════════
# FUNÇÕES PARA RENOVAÇÃO DE CONTRATOS - DEV_21
# ════════════════════════════════════════════════════════════════════

def preparar_contexto_renovacao(renovacao):
    """
    Prepara contexto para contrato de renovação.
    Adiciona variáveis específicas de renovação ao contexto padrão.
    """
    # Usar contexto padrão da nova locação
    contexto = preparar_contexto_contrato(renovacao.nova_locacao)
    
    # ✅ ADICIONAR VARIÁVEIS DE RENOVAÇÃO
    contexto['eh_renovacao'] = True
    contexto['contrato_anterior'] = renovacao.locacao_original.numero_contrato
    contexto['vigencia_anterior_inicio'] = formatar_data(renovacao.locacao_original.data_inicio)
    contexto['vigencia_anterior_fim'] = formatar_data(renovacao.locacao_original.data_fim)
    contexto['valor_anterior'] = formatar_moeda(renovacao.locacao_original.valor_aluguel)
    contexto['valor_novo'] = formatar_moeda(renovacao.novo_valor_aluguel)
    contexto['aumento_percentual'] = f"{renovacao.aumento_percentual:.1f}%"
    contexto['diferenca_valor'] = formatar_moeda(renovacao.diferenca_aluguel)
    
    return contexto


def gerar_docx_contrato_renovacao(renovacao):
    """
    Gera DOCX de contrato de renovação.
    Usa o mesmo sistema de templates, mas com variáveis de renovação.
    """
    # Buscar template (mesmo sistema de locações)
    template_obj = buscar_template_contrato(renovacao.nova_locacao)
    
    if not template_obj:
        raise Exception('Nenhum template de contrato encontrado.')
    
    # Preparar contexto COM VARIÁVEIS DE RENOVAÇÃO
    contexto = preparar_contexto_renovacao(renovacao)
    
    # Carregar template
    doc_template = DocxTemplate(template_obj.arquivo_template.path)
    
    # Preencher template
    doc_template.render(contexto)
    
    # Salvar em memória
    docx_io = io.BytesIO()
    doc_template.save(docx_io)
    docx_io.seek(0)
    
    # Adicionar cabeçalho e rodapé personalizados
    doc = Document(docx_io)
    adicionar_cabecalho_rodape(doc, renovacao.nova_locacao.numero_contrato)
    
    # Salvar novamente
    final_io = io.BytesIO()
    doc.save(final_io)
    final_io.seek(0)
    
    return final_io
