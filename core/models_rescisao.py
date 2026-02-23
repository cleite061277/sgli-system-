"""
════════════════════════════════════════════════════════════════════
MÓDULO: RESCISÃO DE CONTRATOS
Arquivo: core/models_rescisao.py
Descrição: Sistema completo de rescisão contratual com logs judiciais
Autor: Habitat Pro - Desenvolvimento
Data: 2026-02-11
════════════════════════════════════════════════════════════════════
"""

import uuid
import hashlib
import json
from decimal import Decimal
from datetime import timedelta

from django.db import models
from django.utils import timezone
from django.core.exceptions import ValidationError
from django.utils.translation import gettext_lazy as _

# Imports do sistema existente
from .models import BaseModel, Locacao


# ════════════════════════════════════════════════════════════════
# CLASSE ABSTRATA: ComandaBase
# Unifica comportamento de Comanda (aluguel) e ComandaRescisao
# ════════════════════════════════════════════════════════════════

class ComandaBase(BaseModel):
    """
    Classe abstrata com comportamento comum de todas as comandas.
    
    Herança:
    - Comanda (comandas mensais de aluguel) 
    - ComandaRescisao (parcelas de rescisão)
    
    Features:
    - Token público UUID para acesso sem login
    - Sistema de notificações (pré e pós vencimento)
    - Controle de status automático
    - Logs de acessos
    """
    
    class StatusComanda(models.TextChoices):
        PENDENTE = 'PENDENTE', 'Pendente'
        PAGO = 'PAGO', 'Pago'
        ATRASADO = 'ATRASADO', 'Atrasado'
        CANCELADO = 'CANCELADO', 'Cancelado'
    
    # ========================================
    # CAMPOS BÁSICOS
    # ========================================
    
    numero_parcela = models.IntegerField(
        verbose_name=_('Número da Parcela'),
        help_text=_('Número sequencial da parcela')
    )
    
    valor = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        verbose_name=_('Valor'),
        help_text=_('Valor total da parcela')
    )
    
    data_vencimento = models.DateField(
        verbose_name=_('Data de Vencimento'),
        help_text=_('Data de vencimento do pagamento')
    )
    
    data_pagamento = models.DateField(
        null=True,
        blank=True,
        verbose_name=_('Data de Pagamento'),
        help_text=_('Data em que foi efetivamente pago')
    )
    
    descricao = models.CharField(
        max_length=200,
        verbose_name=_('Descrição'),
        help_text=_('Descrição da cobrança')
    )
    
    status = models.CharField(
        max_length=20,
        choices=StatusComanda.choices,
        default=StatusComanda.PENDENTE,
        verbose_name=_('Status'),
        help_text=_('Status atual da comanda')
    )
    
    comprovante = models.FileField(
        upload_to='comprovantes/',
        blank=True,
        null=True,
        verbose_name=_('Comprovante'),
        help_text=_('Comprovante de pagamento')
    )
    
    # ========================================
    # PAGAMENTO ONLINE (FUTURO - ASAAS/MERCADO PAGO)
    # ========================================
    
    gateway_payment_id = models.CharField(
        max_length=100,
        blank=True,
        verbose_name=_('ID Gateway'),
        help_text=_('ID do pagamento no gateway (Asaas/Mercado Pago)')
    )
    
    pix_qrcode = models.TextField(
        blank=True,
        verbose_name=_('QR Code PIX'),
        help_text=_('Código PIX para pagamento')
    )
    
    boleto_url = models.URLField(
        blank=True,
        verbose_name=_('URL Boleto'),
        help_text=_('Link para boleto bancário')
    )
    
    # ========================================
    # CONTROLE DE NOTIFICAÇÕES
    # ========================================
    
    notificacao_enviada = models.BooleanField(
        default=False,
        verbose_name=_('Notificação Enviada'),
        help_text=_('Se notificação de cobrança foi enviada')
    )
    
    data_notificacao = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name=_('Data Notificação'),
        help_text=_('Quando foi enviada a notificação')
    )
    
    notificacao_atraso_enviada = models.BooleanField(
        default=False,
        verbose_name=_('Notificação Atraso Enviada'),
        help_text=_('Se notificação de atraso foi enviada')
    )
    
    data_notificacao_atraso = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name=_('Data Notificação Atraso'),
        help_text=_('Quando foi enviada notificação de atraso')
    )
    
    # ========================================
    # TOKEN PÚBLICO (SEM AUTENTICAÇÃO)
    # ========================================
    
    token_publico = models.UUIDField(
        default=uuid.uuid4,
        unique=True,
        db_index=True,
        editable=False,
        verbose_name=_('Token Público'),
        help_text=_('Token UUID para acesso público sem login')
    )
    
    token_expira_em = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name=_('Token Expira Em'),
        help_text=_('Data de expiração do token público')
    )
    
    token_acessos = models.IntegerField(
        default=0,
        verbose_name=_('Acessos ao Token'),
        help_text=_('Quantidade de vezes que o token foi acessado')
    )
    
    token_ultimo_acesso = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name=_('Último Acesso Token'),
        help_text=_('Data/hora do último acesso ao token')
    )
    
    class Meta:
        abstract = True
    
    # ========================================
    # PROPERTIES
    # ========================================
    
    @property
    def esta_atrasada(self):
        """Verifica se comanda está atrasada"""
        if self.status == self.StatusComanda.PAGO:
            return False
        return self.data_vencimento < timezone.now().date()
    
    @property
    def dias_para_vencimento(self):
        """Retorna dias até vencimento (negativo se atrasado)"""
        delta = self.data_vencimento - timezone.now().date()
        return delta.days
    
    @property
    def dias_atraso(self):
        """Retorna dias de atraso (0 se não atrasado)"""
        if not self.esta_atrasada:
            return 0
        delta = timezone.now().date() - self.data_vencimento
        return delta.days
    
    @property
    def token_esta_valido(self):
        """
        Verifica se token ainda é válido.
        SEGURANÇA: token sem expiração definida é considerado INVÁLIDO.
        Expiração deve ser definida explicitamente via renovar_token().
        """
        if not self.token_expira_em:
            return False  # Sem expiração = inválido (seguro por padrão)
        return timezone.now() < self.token_expira_em
    
    @property
    def valor_formatado(self):
        """Retorna valor formatado em R$"""
        return f"R$ {self.valor:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    
    # ========================================
    # MÉTODOS
    # ========================================
    
    def atualizar_status(self):
        """Atualiza status automaticamente baseado em datas"""
        if self.data_pagamento:
            self.status = self.StatusComanda.PAGO
        elif self.esta_atrasada:
            self.status = self.StatusComanda.ATRASADO
        else:
            self.status = self.StatusComanda.PENDENTE
        self.save()
    
    def registrar_pagamento(self, data_pagamento, comprovante=None):
        """
        Registra pagamento da comanda.
        
        Args:
            data_pagamento: Data do pagamento
            comprovante: Arquivo do comprovante (opcional)
        """
        self.data_pagamento = data_pagamento
        self.status = self.StatusComanda.PAGO
        if comprovante:
            self.comprovante = comprovante
        self.save()
        
        # Envia recibo automaticamente
        try:
            self.enviar_recibo()
        except NotImplementedError:
            pass  # Subclasse não implementou ainda
    
    def renovar_token(self, dias_validade=30):
        """
        Renova token público com nova data de expiração.
        dias_validade: padrão 30 dias, máximo 90 dias.
        """
        dias_validade = min(dias_validade, 90)  # Máximo 90 dias
        self.token_expira_em = timezone.now() + timedelta(days=dias_validade)
        self.save(update_fields=['token_expira_em'])
    
    def registrar_acesso_token(self, ip_address=None, user_agent=None):
        """Registra acesso ao token público"""
        self.token_acessos += 1
        self.token_ultimo_acesso = timezone.now()
        self.save(update_fields=['token_acessos', 'token_ultimo_acesso'])
    
    # ========================================
    # MÉTODOS ABSTRATOS (implementar nas subclasses)
    # ========================================
    
    def enviar_notificacao_cobranca(self):
        """Envia notificação de cobrança - implementar na subclasse"""
        raise NotImplementedError("Subclasse deve implementar enviar_notificacao_cobranca()")
    
    def enviar_notificacao_atraso(self):
        """Envia notificação de atraso - implementar na subclasse"""
        raise NotImplementedError("Subclasse deve implementar enviar_notificacao_atraso()")
    
    def enviar_recibo(self):
        """Envia recibo de pagamento - implementar na subclasse"""
        raise NotImplementedError("Subclasse deve implementar enviar_recibo()")
    
    def get_destinatario(self):
        """Retorna destinatário (Locatario) - implementar na subclasse"""
        raise NotImplementedError("Subclasse deve implementar get_destinatario()")


# ════════════════════════════════════════════════════════════════
# MODEL: RescisaoContrato
# ════════════════════════════════════════════════════════════════

class RescisaoContrato(BaseModel):
    """
    Registra processo completo de rescisão de contrato de locação.
    
    Features:
    - Cálculo automático de valores (multa, contas, caução)
    - Geração automática de documentos (Termo de Recebimento de Chaves)
    - Sistema de confirmação via token público
    - Logs completos para uso judicial
    - Vistoria de saída
    - Parcelamento de débitos
    """
    
    class TipoRescisao(models.TextChoices):
        LOCATARIO = 'LOCATARIO', 'Por Iniciativa do Locatário'
        LOCADOR = 'LOCADOR', 'Por Iniciativa do Locador'
        CONSENSUAL = 'CONSENSUAL', 'Consensual'
        JUDICIAL = 'JUDICIAL', 'Judicial'
    
    class MotivoRescisao(models.TextChoices):
        INADIMPLENCIA = 'INADIMPLENCIA', 'Inadimplência'
        DESOCUPACAO = 'DESOCUPACAO', 'Desocupação Voluntária'
        DESCUMPRIMENTO = 'DESCUMPRIMENTO', 'Descumprimento Contratual'
        FIM_PRAZO = 'FIM_PRAZO', 'Fim do Prazo Determinado'
        OUTRO = 'OUTRO', 'Outro Motivo'
    
    class StatusRescisao(models.TextChoices):
        SOLICITADA = 'SOLICITADA', '📝 Solicitada'
        EM_ANALISE = 'EM_ANALISE', '🔍 Em Análise'
        APROVADA = 'APROVADA', '✅ Aprovada'
        VISTORIA_AGENDADA = 'VISTORIA_AGENDADA', '📅 Vistoria Agendada'
        VISTORIA_REALIZADA = 'VISTORIA_REALIZADA', '🏠 Vistoria Realizada'
        DOCUMENTOS_GERADOS = 'DOCUMENTOS_GERADOS', '📄 Documentos Gerados'
        AGUARDANDO_PAGAMENTO = 'AGUARDANDO_PAGAMENTO', '⏳ Aguardando Pagamento'
        FINALIZADA = 'FINALIZADA', '🏁 Finalizada'
        CANCELADA = 'CANCELADA', '❌ Cancelada'
    
    # ========================================
    # RELACIONAMENTOS
    # ========================================
    
    locacao = models.ForeignKey(
        'Locacao',
        on_delete=models.PROTECT,
        related_name='rescisoes',
        verbose_name=_('Locação'),
        help_text=_('Contrato de locação que está sendo rescindido')
    )
    
    solicitante = models.ForeignKey(
        'Usuario',
        on_delete=models.PROTECT,
        related_name='rescisoes_solicitadas',
        verbose_name=_('Solicitante'),
        help_text=_('Usuário do sistema que registrou a solicitação')
    )
    
    # ========================================
    # DADOS DA RESCISÃO
    # ========================================
    
    tipo = models.CharField(
        max_length=20,
        choices=TipoRescisao.choices,
        verbose_name=_('Tipo de Rescisão'),
        help_text=_('Quem está solicitando a rescisão')
    )
    
    motivo = models.CharField(
        max_length=20,
        choices=MotivoRescisao.choices,
        verbose_name=_('Motivo'),
        help_text=_('Motivo da rescisão')
    )
    
    data_solicitacao = models.DateField(
        auto_now_add=True,
        verbose_name=_('Data de Solicitação'),
        help_text=_('Data em que foi solicitada a rescisão')
    )
    
    data_desocupacao = models.DateField(
        verbose_name=_('Data de Desocupação'),
        help_text=_('Data prevista/efetiva de desocupação do imóvel')
    )
    
    data_vistoria = models.DateField(
        null=True,
        blank=True,
        verbose_name=_('Data da Vistoria'),
        help_text=_('Data em que foi realizada a vistoria de saída')
    )
    
    status = models.CharField(
        max_length=30,
        choices=StatusRescisao.choices,
        default=StatusRescisao.SOLICITADA,
        verbose_name=_('Status'),
        help_text=_('Status atual do processo de rescisão')
    )
    
    # ========================================
    # OBSERVAÇÕES
    # ========================================
    
    observacoes_solicitacao = models.TextField(
        blank=True,
        verbose_name=_('Observações da Solicitação'),
        help_text=_('Observações informadas ao solicitar rescisão')
    )
    
    observacoes_vistoria = models.TextField(
        blank=True,
        verbose_name=_('Observações da Vistoria'),
        help_text=_('Observações da vistoria de saída')
    )
    
    # ========================================
    # VISTORIA
    # ========================================
    
    imovel_conforme = models.BooleanField(
        null=True,
        blank=True,
        verbose_name=_('Imóvel Conforme'),
        help_text=_('Se imóvel está nas mesmas condições da entrada')
    )
    
    avarias_detectadas = models.TextField(
        blank=True,
        verbose_name=_('Avarias Detectadas'),
        help_text=_('Descrição detalhada das avarias encontradas')
    )
    
    custo_reparos = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=Decimal('0.00'),
        verbose_name=_('Custo de Reparos'),
        help_text=_('Valor estimado para reparos necessários')
    )
    
    # ========================================
    # VALORES (calculados automaticamente)
    # ========================================
    
    multa_rescisoria = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=Decimal('0.00'),
        verbose_name=_('Multa Rescisória'),
        help_text=_('Multa por rescisão antecipada (ex: 3 aluguéis)')
    )
    
    valor_aluguel_proporcional = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=Decimal('0.00'),
        verbose_name=_('Aluguel Proporcional'),
        help_text=_('Aluguel proporcional aos dias ocupados')
    )
    
    valor_agua = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=Decimal('0.00'),
        verbose_name=_('Valor Água'),
        help_text=_('Conta de água pendente')
    )
    
    valor_luz = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=Decimal('0.00'),
        verbose_name=_('Valor Luz'),
        help_text=_('Conta de luz pendente')
    )
    
    valor_gas = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=Decimal('0.00'),
        verbose_name=_('Valor Gás'),
        help_text=_('Conta de gás pendente')
    )
    
    valor_condominio = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=Decimal('0.00'),
        verbose_name=_('Valor Condomínio'),
        help_text=_('Condomínio pendente')
    )
    
    outras_despesas = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=Decimal('0.00'),
        verbose_name=_('Outras Despesas'),
        help_text=_('Outras despesas pendentes')
    )
    
    valor_caucao = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=Decimal('0.00'),
        verbose_name=_('Valor Caução (deduzir)'),
        help_text=_('Valor da caução a ser deduzido do total devido')
    )
    
    # TOTAIS
    valor_total_devido = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=Decimal('0.00'),
        verbose_name=_('Total Devido'),
        help_text=_('Valor total devido na rescisão')
    )
    
    valor_total_pago = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=Decimal('0.00'),
        verbose_name=_('Total Pago'),
        help_text=_('Valor total já pago')
    )
    
    # ========================================
    # PARCELAMENTO
    # ========================================
    
    quantidade_parcelas = models.IntegerField(
        default=1,
        verbose_name=_('Quantidade de Parcelas'),
        help_text=_('Número de parcelas acordado')
    )
    
    valor_parcela = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=Decimal('0.00'),
        verbose_name=_('Valor da Parcela'),
        help_text=_('Valor de cada parcela')
    )
    
    # ========================================
    # DOCUMENTOS GERADOS
    # ========================================
    
    termo_recebimento_chaves = models.FileField(
        upload_to='rescisoes/termos/',
        blank=True,
        null=True,
        verbose_name=_('Termo de Recebimento de Chaves'),
        help_text=_('Documento Word gerado automaticamente')
    )
    
    laudo_vistoria = models.FileField(
        upload_to='rescisoes/vistorias/',
        blank=True,
        null=True,
        verbose_name=_('Laudo de Vistoria'),
        help_text=_('Laudo da vistoria de saída')
    )
    
    # ========================================
    # LOGS E COMPROVANTES (USO JUDICIAL)
    # ========================================
    
    log_solicitacao = models.JSONField(
        default=dict,
        blank=True,
        verbose_name=_('Log de Solicitação'),
        help_text=_('Log detalhado da solicitação inicial')
    )
    
    log_confirmacao_inquilino = models.JSONField(
        default=dict,
        blank=True,
        verbose_name=_('Log de Confirmação do Inquilino'),
        help_text=_('Log da confirmação via link público')
    )
    
    log_vistoria = models.JSONField(
        default=dict,
        blank=True,
        verbose_name=_('Log de Vistoria'),
        help_text=_('Registro completo da vistoria')
    )
    
    log_comunicacoes = models.JSONField(
        default=list,
        blank=True,
        verbose_name=_('Log de Comunicações'),
        help_text=_('Histórico completo de emails/WhatsApp enviados')
    )
    
    # ========================================
    # TOKENS PARA CONFIRMAÇÃO
    # ========================================
    
    token_confirmacao = models.UUIDField(
        default=uuid.uuid4,
        unique=True,
        db_index=True,
        editable=False,
        verbose_name=_('Token de Confirmação'),
        help_text=_('Token para inquilino confirmar rescisão')
    )
    
    token_confirmacao_expira = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name=_('Token Expira Em'),
        help_text=_('Data de expiração do token (7 dias padrão)')
    )
    
    data_confirmacao_inquilino = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name=_('Data Confirmação Inquilino'),
        help_text=_('Quando inquilino confirmou via link')
    )
    
    confirmado_por_inquilino = models.BooleanField(
        default=False,
        verbose_name=_('Confirmado pelo Inquilino'),
        help_text=_('Se inquilino confirmou a rescisão')
    )
    
    class Meta:
        db_table = 'core_rescisao_contrato'
        verbose_name = 'Rescisão de Contrato'
        verbose_name_plural = '🔄 3.1 Rescisões de Contratos'
        ordering = ['-data_solicitacao']
        indexes = [
            models.Index(fields=['status']),
            models.Index(fields=['data_desocupacao']),
            models.Index(fields=['tipo']),
            models.Index(fields=['token_confirmacao']),
        ]
    

    
    # ========================================
    # PROPERTIES
    # ========================================

    @property
    def saldo_devedor(self):
        """Saldo devedor da rescisão (total_devido - total_pago)."""
        return self.valor_total_devido - self.valor_total_pago

    @property
    def locatario(self):
        """Atalho para acessar o locatário via FK locacao."""
        return self.locacao.locatario if self.locacao_id else None

    @property
    def imovel(self):
        """Atalho para acessar o imóvel via FK locacao."""
        return self.locacao.imovel if self.locacao_id else None

    @property
    def numero_contrato(self):
        """Número do contrato de locação."""
        return self.locacao.numero_contrato if self.locacao_id else 'N/A'

    @property
    def locatario_nome(self):
        """Nome do locatário (para exibição rápida)."""
        return self.locacao.locatario.nome_razao_social if self.locacao_id else 'N/A'

    @property
    def imovel_endereco(self):
        """Endereço do imóvel (para exibição rápida)."""
        if self.locacao_id and self.locacao.imovel:
            return self.imovel.endereco
        return 'N/A'

    def __str__(self):
        return f"Rescisão #{str(self.id)[:8]} - {self.locacao.numero_contrato} - {self.get_status_display()}"
    
    # ========================================
    # SAVE OVERRIDE
    # ========================================
    
    def save(self, *args, **kwargs):
        """Override save para gerar token de confirmação automaticamente"""
        if not self.pk and not self.token_confirmacao_expira:
            # Primeiro save: gera token com 7 dias de validade
            self.token_confirmacao_expira = timezone.now() + timedelta(days=7)
        
        super().save(*args, **kwargs)
        
        # Atualiza status do contrato de locação
        if self.status == self.StatusRescisao.FINALIZADA:
            self.locacao.status = 'RESCINDED'
            self.locacao.data_rescisao = self.data_desocupacao
            self.locacao.save(update_fields=['status', 'data_rescisao'])
    
    # ========================================
    # MÉTODOS DE CÁLCULO
    # ========================================
    
    def calcular_valores(self):
        """
        Calcula automaticamente TODOS os valores da rescisão.
        
        Regras:
        - Multa rescisória: 3 aluguéis se locatário rescinde
        - Aluguel proporcional: dias ocupados no mês
        - Total = multa + proporcional + contas - caução
        """
        contrato = self.locacao
        
        # 1. MULTA RESCISÓRIA
        if self.tipo == self.TipoRescisao.LOCATARIO:
            self.multa_rescisoria = contrato.valor_aluguel * Decimal('3.00')
        else:
            self.multa_rescisoria = Decimal('0.00')
        
        # 2. ALUGUEL PROPORCIONAL
        # TODO: Implementar cálculo preciso baseado em comandas pagas
        # Por ora, mantém valor manual
        
        # 3. TOTAL DEVIDO
        self.valor_total_devido = (
            self.multa_rescisoria +
            self.valor_aluguel_proporcional +
            self.valor_agua +
            self.valor_luz +
            self.valor_gas +
            self.valor_condominio +
            self.custo_reparos +
            self.outras_despesas -
            self.valor_caucao
        )
        
        # Garantir que não fique negativo
        if self.valor_total_devido < Decimal('0.00'):
            self.valor_total_devido = Decimal('0.00')
        
        # 4. VALOR DA PARCELA
        if self.quantidade_parcelas > 0:
            self.valor_parcela = self.valor_total_devido / Decimal(str(self.quantidade_parcelas))
        else:
            self.valor_parcela = self.valor_total_devido
        
        self.save()
    
    def gerar_comandas_rescisao(self):
        """
        Gera comandas de pagamento vinculadas à rescisão.
        Remove comandas antigas se existirem.
        """
        from datetime import timedelta
        
        # Limpa comandas antigas
        ComandaRescisao.objects.filter(rescisao=self).delete()
        
        # Data base: 10 dias após desocupação para primeiro vencimento
        data_vencimento_base = self.data_desocupacao + timedelta(days=10)
        
        # Gera novas comandas
        for i in range(self.quantidade_parcelas):
            ComandaRescisao.objects.create(
                rescisao=self,
                numero_parcela=i + 1,
                valor=self.valor_parcela,
                data_vencimento=data_vencimento_base + timedelta(days=30 * i),
                descricao=f"Parcela {i+1}/{self.quantidade_parcelas} - Rescisão Contrato {self.locacao.numero_contrato}",
                status='PENDENTE'
            )
        
        # Atualiza status
        if self.status == self.StatusRescisao.DOCUMENTOS_GERADOS:
            self.status = self.StatusRescisao.AGUARDANDO_PAGAMENTO
            self.save()
    
    def gerar_termo_recebimento_chaves(self):
        """
        Gera documento Word do Termo de Recebimento de Chaves.
        Baseado no template fornecido.
        """
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        from django.core.files import File
        
        doc = Document()
        
        # TÍTULO
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run('TERMO DE RECEBIMENTO DE CHAVES')
        run.bold = True
        run.underline = True
        run.font.size = Pt(14)
        
        doc.add_paragraph()  # Espaço
        
        # DADOS
        locatario = self.locacao.locatario
        locador = self.locacao.imovel.locador
        imovel = self.locacao.imovel
        
        # Parágrafo inicial
        doc.add_paragraph(
            f"Por este particular instrumento, de um lado "
            f"{locatario.nome_razao_social} (RG: {locatario.rg or 'N/A'} e CPF: {locatario.cpf_cnpj}) "
            f"e de outro {locador.nome_razao_social}, "
            f"já qualificados respectivamente como LOCATÁRIO e LOCADOR, "
            f"no Contrato de Locação firmado entre partes em data de "
            f"{self.locacao.data_inicio.strftime('%d/%m/%Y')} tendo por objeto "
            f"o imóvel situado na {imovel.endereco}, {imovel.numero}, {imovel.bairro}, "
            f"{imovel.cidade}-{imovel.estado}, CEP: {imovel.cep}; resolve o locatário, "
            f"por decisão unilateral, rescindir na data de "
            f"{self.data_desocupacao.strftime('%d/%m/%Y')} o contrato de locação."
        )
        
        # Vistoria
        doc.add_paragraph(
            f"Os Locadores informam terem recebido as chaves do imóvel. "
            f"Quanto à vistoria, fora realizada em "
            f"{self.data_vistoria.strftime('%d/%m/%Y') if self.data_vistoria else '___/___/___'}, "
            f"estando o imóvel nas seguintes condições:"
        )
        
        # Condições
        p1 = doc.add_paragraph()
        p1.add_run(f"1 ({'X' if self.imovel_conforme else ' '}) Nas mesmas condições da assinatura do contrato.")
        
        p2 = doc.add_paragraph()
        p2.add_run(f"2 ({' ' if self.imovel_conforme else 'X'}) Não está de acordo com as condições.")
        
        if not self.imovel_conforme and self.avarias_detectadas:
            doc.add_paragraph(f"\nAvarias: {self.avarias_detectadas}")
            doc.add_paragraph(
                "2.1 O Locatário reconhece tal fato e assume total responsabilidade "
                "pelos custos de restituição do imóvel."
            )
        
        # Valores
        doc.add_paragraph(
            f"\nConsiderando a decisão unilateral, aplica-se a multa contratual de três aluguéis, "
            f"perfazendo R$ {self.multa_rescisoria:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') + "."
        )
        
        # Detalhamento
        detalhes = []
        if self.valor_agua > 0:
            detalhes.append(f"água (R$ {self.valor_agua:,.2f})".replace(',', 'X').replace('.', ',').replace('X', '.'))
        if self.valor_luz > 0:
            detalhes.append(f"luz (R$ {self.valor_luz:,.2f})".replace(',', 'X').replace('.', ',').replace('X', '.'))
        if self.valor_condominio > 0:
            detalhes.append(f"condomínio (R$ {self.valor_condominio:,.2f})".replace(',', 'X').replace('.', ',').replace('X', '.'))
        
        if detalhes:
            doc.add_paragraph(f"Despesas: {', '.join(detalhes)}.")
        
        total_fmt = f"R$ {self.valor_total_devido:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        doc.add_paragraph(f"Total da dívida: {total_fmt}.")
        
        # Compromisso
        comp = doc.add_paragraph()
        comp.add_run(
            "\nLocatário reconhece os valores e assume obrigação da quitação conforme plano abaixo:"
        ).bold = True
        
        # Parcelas
        parcela_fmt = f"R$ {self.valor_parcela:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        doc.add_paragraph(f"a) {self.quantidade_parcelas} Parcela(s) de {parcela_fmt}")
        
        # Multa atraso
        multa_p = doc.add_paragraph()
        multa_p.add_run(
            "O atraso incidirá multa de 10% + correção monetária mensal (IPCA) "
            "conforme cláusulas contratuais."
        ).underline = True
        
        # Responsabilidades
        doc.add_paragraph(
            "\nTodas as despesas do imóvel são de responsabilidade do LOCATÁRIO."
        )
        
        doc.add_paragraph(
            "Para efeitos legais, o Locatário reconhece e obriga-se ao pagamento "
            "da dívida em sua integralidade."
        )
        
        # Data e assinaturas
        doc.add_paragraph(f"\n\nParanaguá, {self.data_desocupacao.strftime('%d/%m/%Y')}\n")
        
        assin = doc.add_paragraph()
        assin.add_run(f"Locatário: {locatario.nome_razao_social}     |     ")
        assin.add_run(f"Locador: {locador.nome_razao_social}")
        
        # Salvar
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f'termo_rescisao_{str(self.id)[:8]}_{self.locacao.numero_contrato}.docx'
        self.termo_recebimento_chaves.save(filename, File(buffer), save=True)
        
        # Atualiza status
        if self.status == self.StatusRescisao.VISTORIA_REALIZADA:
            self.status = self.StatusRescisao.DOCUMENTOS_GERADOS
            self.save()
        
        return self.termo_recebimento_chaves.url
    
    # ========================================
    # MÉTODOS DE LOG (USO JUDICIAL)
    # ========================================
    
    def registrar_solicitacao(self, usuario_admin, origem='manual', observacoes='', email_anexo=None):
        """Registra solicitação inicial com todos os detalhes"""
        self.log_solicitacao = {
            'data': timezone.now().isoformat(),
            'usuario_admin': usuario_admin.email,
            'usuario_admin_nome': usuario_admin.get_full_name(),
            'origem': origem,
            'email_original': email_anexo,
            'observacoes': observacoes,
        }
        self.save(update_fields=['log_solicitacao'])
    
    def gerar_token_confirmacao(self, dias_validade=7):
        """Gera token novo para confirmação"""
        self.token_confirmacao = uuid.uuid4()
        self.token_confirmacao_expira = timezone.now() + timedelta(days=dias_validade)
        self.save(update_fields=['token_confirmacao', 'token_confirmacao_expira'])
        return self.token_confirmacao
    
    def registrar_confirmacao(self, request):
        """Registra confirmação do inquilino via link público"""
        dados = {
            'data_acesso': timezone.now().isoformat(),
            'ip_acesso': self._get_client_ip(request),
            'user_agent': request.META.get('HTTP_USER_AGENT', 'unknown'),
            'confirmado': True,
            'timestamp_confirmacao': timezone.now().isoformat(),
            'termos_aceitos': True,
            'assinatura_eletronica': f"{self.locacao.locatario.nome_razao_social} - {timezone.now().strftime('%d/%m/%Y %H:%M')}",
        }
        
        # Hash para integridade
        dados_str = json.dumps(dados, sort_keys=True)
        dados['hash_confirmacao'] = f"sha256:{hashlib.sha256(dados_str.encode()).hexdigest()}"
        
        self.log_confirmacao_inquilino = dados
        self.confirmado_por_inquilino = True
        self.data_confirmacao_inquilino = timezone.now()
        self.status = self.StatusRescisao.APROVADA
        self.save()
        
        # Registra comunicação
        self.registrar_comunicacao('confirmacao_web', self.locacao.locatario.email, True, 'Confirmação via link público')
    
    def registrar_comunicacao(self, tipo, destinatario, sucesso=True, detalhes='', message_id=None):
        """Registra todas as comunicações enviadas"""
        if not isinstance(self.log_comunicacoes, list):
            self.log_comunicacoes = []
        
        self.log_comunicacoes.append({
            'data': timezone.now().isoformat(),
            'tipo': tipo,
            'destinatario': destinatario,
            'status': 'enviado' if sucesso else 'erro',
            'detalhes': detalhes,
            'message_id': message_id,
        })
        self.save(update_fields=['log_comunicacoes'])
    
    @staticmethod
    def _get_client_ip(request=None):
        """Obtém IP real do cliente"""
        if not request:
            return 'unknown'
        
        x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
        if x_forwarded_for:
            ip = x_forwarded_for.split(',')[0]
        else:
            ip = request.META.get('REMOTE_ADDR')
        return ip
    
    # ========================================
    # PROPERTIES
    # ========================================
    
    @property
    def saldo_devedor(self):
        """Retorna saldo ainda em aberto"""
        return self.valor_total_devido - self.valor_total_pago
    
    @property
    def percentual_pago(self):
        """Retorna percentual pago (0-100)"""
        if self.valor_total_devido == 0:
            return Decimal('100.00')
        return (self.valor_total_pago / self.valor_total_devido) * Decimal('100.00')
    
    @property
    def token_confirmacao_valido(self):
        """Verifica se token ainda é válido"""
        if not self.token_confirmacao_expira:
            return True
        return timezone.now() < self.token_confirmacao_expira


# ════════════════════════════════════════════════════════════════
# MODEL: ComandaRescisao
# ════════════════════════════════════════════════════════════════

class ComandaRescisao(ComandaBase):
    """
    Comandas de pagamento vinculadas à rescisão de contrato.
    Funciona exatamente como Comanda (aluguel), mas para rescisão.
    """
    
    rescisao = models.ForeignKey(
        RescisaoContrato,
        on_delete=models.CASCADE,
        related_name='comandas',
        verbose_name=_('Rescisão'),
        help_text=_('Rescisão de contrato associada')
    )

    valor_pago = models.DecimalField(
        _('Valor Pago'),
        max_digits=10,
        decimal_places=2,
        default=0,
        help_text=_('Soma de todos os pagamentos confirmados desta comanda')
    )

    class Meta:
        db_table = 'core_comanda_rescisao'
        verbose_name = 'Comanda de Rescisão'
        verbose_name_plural = '🔄 3.2 Comandas de Rescisões'
        ordering = ['numero_parcela']
        indexes = [
            models.Index(fields=['status']),
            models.Index(fields=['data_vencimento']),
            models.Index(fields=['token_publico']),
        ]
    

    
    @property
    def saldo(self):
        """Saldo: valor_pago - valor (negativo = deve, positivo = crédito)."""
        from decimal import Decimal
        if not hasattr(self, 'valor_pago'):
            return Decimal('0.00')
        return self.valor_pago - self.valor


    
    # ========================================
    # PROPERTIES DE NAVEGAÇÃO (Atalhos FK)
    # ========================================
    
    @property
    def locacao(self):
        """Atalho para acessar a locação (contrato original)."""
        return self.rescisao.locacao if self.rescisao else None
    
    @property
    def locatario(self):
        """Atalho para acessar o locatário."""
        if self.rescisao and self.rescisao.locacao:
            return self.rescisao.locacao.locatario
        return None
    
    @property
    def imovel(self):
        """Atalho para acessar o imóvel."""
        if self.rescisao and self.rescisao.locacao:
            return self.rescisao.locacao.imovel
        return None
    
    @property
    def numero_contrato(self):
        """Número do contrato de locação."""
        if self.rescisao and self.rescisao.locacao:
            return self.rescisao.locacao.numero_contrato
        return 'N/A'
    
    @property
    def locatario_nome(self):
        """Nome do locatário (para exibição rápida)."""
        if self.locatario:
            return self.locatario.nome_razao_social
        return 'N/A'
    
    @property
    def imovel_endereco(self):
        """Endereço do imóvel (para exibição rápida)."""
        if self.imovel:
            return self.imovel.endereco
        return 'N/A'

    def __str__(self):
        return f"Rescisão {str(self.rescisao.id)[:8]} - Parcela {self.numero_parcela}/{self.rescisao.quantidade_parcelas}"
    
    def save(self, *args, **kwargs):
        """Override save para atualizar total_pago da rescisão"""
        is_new = self.pk is None
        super().save(*args, **kwargs)
        
        # Se mudou para PAGO, atualizar total_pago da rescisão
        if not is_new and self.status == self.StatusComanda.PAGO:
            self.rescisao.valor_total_pago = sum(
                c.valor for c in self.rescisao.comandas.filter(status='PAGO')
            )
            self.rescisao.save(update_fields=['valor_total_pago'])
            
            # Se todas pagas, finalizar rescisão
            total_comandas = self.rescisao.comandas.count()
            comandas_pagas = self.rescisao.comandas.filter(status='PAGO').count()
            if total_comandas == comandas_pagas:
                self.rescisao.status = 'FINALIZADA'
                self.rescisao.save(update_fields=['status'])
    
    # ========================================
    # IMPLEMENTAÇÃO DOS MÉTODOS ABSTRATOS
    # ========================================
    
    def get_destinatario(self):
        """Retorna o locatário"""
        return self.rescisao.locacao.locatario
    
    def enviar_notificacao_cobranca(self):
        """Envia email de cobrança"""
        # TODO: Implementar integração com EmailService
        try:
            from core.services.email_service import EmailService
            EmailService.enviar_comanda_rescisao(self)
            self.notificacao_enviada = True
            self.data_notificacao = timezone.now()
            self.save(update_fields=['notificacao_enviada', 'data_notificacao'])
            return True
        except Exception as e:
            print(f"Erro ao enviar notificação: {e}")
            return False
    
    def enviar_notificacao_atraso(self):
        """Envia email de atraso"""
        try:
            from core.services.email_service import EmailService
            EmailService.enviar_atraso_rescisao(self)
            self.notificacao_atraso_enviada = True
            self.data_notificacao_atraso = timezone.now()
            self.save(update_fields=['notificacao_atraso_enviada', 'data_notificacao_atraso'])
            return True
        except Exception as e:
            print(f"Erro ao enviar notificação de atraso: {e}")
            return False
    
    def enviar_recibo(self):
        """Envia recibo de pagamento"""
        try:
            from core.services.email_service import EmailService
            EmailService.enviar_recibo_rescisao(self)
            return True
        except Exception as e:
            print(f"Erro ao enviar recibo: {e}")
            return False
    
    def gerar_link_whatsapp(self):
        """Gera link wa.me para WhatsApp"""
        import urllib.parse
        
        destinatario = self.get_destinatario()
        telefone = destinatario.telefone.replace('(', '').replace(')', '').replace('-', '').replace(' ', '').replace('+', '')
        
        if not telefone.startswith('55'):
            telefone = '55' + telefone
        
        # Mensagem formatada
        mensagem = (
            f"🏠 *RESCISÃO DE CONTRATO* 🏠\n\n"
            f"Olá, {destinatario.nome_razao_social}! 👋\n\n"
            f"📋 *Detalhes da Parcela:*\n"
            f"━━━━━━━━━━━━━━━━━\n"
            f"📌 Parcela: {self.numero_parcela}/{self.rescisao.quantidade_parcelas}\n"
            f"💰 Valor: {self.valor_formatado}\n"
            f"📅 Vencimento: {self.data_vencimento.strftime('%d/%m/%Y')}\n\n"
        )
        
        if self.esta_atrasada:
            mensagem += f"⚠️ *PAGAMENTO EM ATRASO*\n"
            mensagem += f"🔴 Atraso: {self.dias_atraso} dia(s)\n\n"
        else:
            mensagem += f"✅ *No prazo*\n\n"
        
        # Link visualização
        from django.conf import settings
        link = f"{settings.SITE_URL}/comanda-rescisao/{self.token_publico}/"
        mensagem += f"🔗 Visualizar:\n{link}\n\n"
        mensagem += "Dúvidas? Estamos à disposição! 💙"
        
        mensagem_encoded = urllib.parse.quote(mensagem)
        return f"https://wa.me/{telefone}?text={mensagem_encoded}"
