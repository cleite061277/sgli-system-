"""
Gerador de contratos PDF e DOCX - CORRIGIDO
"""

from django.http import HttpResponse
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from io import BytesIO


def gerar_contrato_pdf(locacao):
    """Gera contrato PDF"""
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4, 
        rightMargin=2*cm, 
        leftMargin=2*cm, 
        topMargin=2*cm, 
        bottomMargin=2*cm
    )
    elements = []
    styles = getSampleStyleSheet()
    
    # Título
    titulo_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#2c3e50'),
        spaceAfter=30,
        alignment=1,
        fontName='Helvetica-Bold'
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6,
        alignment=4
    )
    
    titulo = Paragraph("CONTRATO DE LOCACAO DE IMOVEL", titulo_style)
    elements.append(titulo)
    elements.append(Spacer(1, 0.5*cm))
    
    # Info básica
    elements.append(Paragraph("<b>Numero do Contrato:</b> {}".format(locacao.numero_contrato), normal_style))
    elements.append(Paragraph("<b>Data:</b> {}".format(datetime.now().strftime('%d/%m/%Y')), normal_style))
    elements.append(Paragraph("<b>Vigencia:</b> {} ate {}".format(
        locacao.data_inicio.strftime('%d/%m/%Y'),
        locacao.data_fim.strftime('%d/%m/%Y')
    ), normal_style))
    elements.append(Spacer(1, 0.5*cm))
    
    # Locador
    elements.append(Paragraph("<b>LOCADOR:</b> {}".format(locacao.imovel.locador.nome_razao_social), normal_style))
    elements.append(Paragraph("<b>CPF/CNPJ:</b> {}".format(locacao.imovel.locador.cpf_cnpj), normal_style))
    elements.append(Spacer(1, 0.3*cm))
    
    # Locatário
    elements.append(Paragraph("<b>LOCATARIO:</b> {}".format(locacao.locatario.nome_razao_social), normal_style))
    elements.append(Paragraph("<b>CPF/CNPJ:</b> {}".format(locacao.locatario.cpf_cnpj), normal_style))
    
    if hasattr(locacao.locatario, 'rg') and locacao.locatario.rg:
        elements.append(Paragraph("<b>RG:</b> {}".format(locacao.locatario.rg), normal_style))
    
    if hasattr(locacao.locatario, 'fiador') and locacao.locatario.fiador:
        elements.append(Spacer(1, 0.3*cm))
        elements.append(Paragraph("<b>FIADOR:</b> {}".format(locacao.locatario.fiador.nome_completo), normal_style))
        elements.append(Paragraph("<b>CPF:</b> {}".format(locacao.locatario.fiador.cpf), normal_style))
    
    elements.append(Spacer(1, 0.5*cm))
    
    # Imóvel
    elements.append(Paragraph("<b>IMOVEL:</b> {}".format(locacao.imovel.endereco_completo), normal_style))
    elements.append(Paragraph("<b>Codigo:</b> {}".format(locacao.imovel.codigo_imovel), normal_style))
    elements.append(Spacer(1, 0.5*cm))
    
    # Valores
    valor_formatado = "{:,.2f}".format(float(locacao.valor_aluguel)).replace(',', 'X').replace('.', ',').replace('X', '.')
    elements.append(Paragraph("<b>Valor do Aluguel:</b> R$ {}".format(valor_formatado), normal_style))
    elements.append(Paragraph("<b>Vencimento:</b> Dia {} de cada mes".format(locacao.dia_vencimento), normal_style))
    elements.append(Spacer(1, 1*cm))
    
    # Assinaturas
    elements.append(Paragraph("_______________, {}".format(datetime.now().strftime('%d/%m/%Y')), normal_style))
    elements.append(Spacer(1, 1*cm))
    
    assinaturas = [
        ['_' * 40, '_' * 40],
        ['LOCADOR', 'LOCATARIO']
    ]
    tabela = Table(assinaturas, colWidths=[8*cm, 8*cm])
    tabela.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 1), (-1, 1), 'Helvetica-Bold')
    ]))
    elements.append(tabela)
    
    doc.build(elements)
    buffer.seek(0)
    
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="Contrato_{}.pdf"'.format(locacao.numero_contrato)
    return response


def gerar_contrato_docx(locacao):
    """Gera contrato DOCX"""
    
    doc = Document()
    
    # Título
    titulo = doc.add_heading('CONTRATO DE LOCACAO DE IMOVEL', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Info
    doc.add_paragraph('Numero: {}'.format(locacao.numero_contrato))
    doc.add_paragraph('Data: {}'.format(datetime.now().strftime("%d/%m/%Y")))
    doc.add_paragraph('Vigencia: {} ate {}'.format(
        locacao.data_inicio.strftime("%d/%m/%Y"),
        locacao.data_fim.strftime("%d/%m/%Y")
    ))
    doc.add_paragraph()
    
    # Partes
    p = doc.add_paragraph()
    p.add_run('LOCADOR: ').bold = True
    p.add_run(locacao.imovel.locador.nome_razao_social)
    
    p = doc.add_paragraph()
    p.add_run('LOCATARIO: ').bold = True
    p.add_run(locacao.locatario.nome_razao_social)
    
    if hasattr(locacao.locatario, 'fiador') and locacao.locatario.fiador:
        p = doc.add_paragraph()
        p.add_run('FIADOR: ').bold = True
        p.add_run(locacao.locatario.fiador.nome_completo)
    
    doc.add_paragraph()
    
    # Imóvel
    p = doc.add_paragraph()
    p.add_run('IMOVEL: ').bold = True
    p.add_run(locacao.imovel.endereco_completo)
    
    doc.add_paragraph()
    
    # Valores
    valor_formatado = "{:,.2f}".format(float(locacao.valor_aluguel)).replace(',', 'X').replace('.', ',').replace('X', '.')
    p = doc.add_paragraph()
    p.add_run('Aluguel: ').bold = True
    p.add_run('R$ {}'.format(valor_formatado))
    
    p = doc.add_paragraph()
    p.add_run('Vencimento: ').bold = True
    p.add_run('Dia {}'.format(locacao.dia_vencimento))
    
    # Assinaturas
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('_______________, {}'.format(datetime.now().strftime("%d/%m/%Y")))
    doc.add_paragraph()
    
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = '_' * 30
    table.rows[0].cells[1].text = '_' * 30
    table.rows[1].cells[0].text = 'LOCADOR'
    table.rows[1].cells[1].text = 'LOCATARIO'
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="Contrato_{}.docx"'.format(locacao.numero_contrato)
    return response
