"""
Sistema de geração de documentos Word para SGLI.
Gera contratos, recibos e relatórios com substituição automática de campos.
"""

import os
from datetime import datetime
from decimal import Decimal
from typing import Dict, Any
from docx import Document
from docxtpl import DocxTemplate
from .odt_converter import processar_template_odt
from django.conf import settings
from .models import Locacao, Comanda, Pagamento
from .utils import formatar_moeda_brasileira

class DocumentGenerator:
    """Gerador principal de documentos."""
    
    def __init__(self):
        self.template_dir = os.path.join(settings.MEDIA_ROOT, 'templates_word')
        self.output_dir = os.path.join(settings.MEDIA_ROOT, 'documentos_gerados')
        
        # Criar diretórios se não existirem
        os.makedirs(self.template_dir, exist_ok=True)
        os.makedirs(self.output_dir, exist_ok=True)
    
    def gerar_contrato_locacao(self, locacao_id: int) -> str:
        """Gerar contrato de locação usando template apropriado."""
        from .models import Locacao, TemplateContrato
        locacao = Locacao.objects.get(id=locacao_id)
        
        # Buscar template apropriado
        template_path = self._selecionar_template_contrato(locacao)
        
        # Detectar tipo de arquivo
        context = self._preparar_contexto_contrato(locacao)
        
        if template_path.endswith('.odt'):
            doc = processar_template_odt(template_path, context)
        else:
            doc = DocxTemplate(template_path)
            doc.render(context)
        
        output_filename = f"contrato_{locacao.numero_contrato}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = os.path.join(self.output_dir, output_filename)
        doc.save(output_path)
        
        return output_filename
    
    def _selecionar_template_contrato(self, locacao):
        """Selecionar o template de contrato mais apropriado."""
        from .models import TemplateContrato
        
        # 1. Tentar template específico do locador
        template = TemplateContrato.objects.filter(
            locador=locacao.imovel.locador,
            is_active=True
        ).first()
        
        if template and template.arquivo_template:
            return template.arquivo_template.path
        
        # 2. Tentar template por tipo de imóvel
        template = TemplateContrato.objects.filter(
            tipo_imovel=locacao.imovel.tipo_imovel,
            is_active=True
        ).first()
        
        if template and template.arquivo_template:
            return template.arquivo_template.path
        
        # 3. Usar template padrão
        template = TemplateContrato.objects.filter(
            is_default=True,
            is_active=True
        ).first()
        
        if template and template.arquivo_template:
            return template.arquivo_template.path
        
        # 4. Criar template básico se não houver nenhum
        default_path = os.path.join(self.template_dir, 'contrato_locacao.docx')
        if not os.path.exists(default_path):
            self._criar_template_contrato(default_path)
        
        return default_path
    
    def gerar_recibo_pagamento(self, pagamento_id: int) -> str:
        """Gerar recibo de pagamento."""
        from .models import Pagamento
        pagamento = Pagamento.objects.get(id=pagamento_id)
        
        template_path = os.path.join(self.template_dir, 'recibo_pagamento.docx')
        
        if not os.path.exists(template_path):
            self._criar_template_recibo(template_path)
        
        doc = DocxTemplate(template_path)
        context = self._preparar_contexto_recibo(pagamento)
        doc.render(context)
        
        output_filename = f"recibo_{pagamento.numero_pagamento}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = os.path.join(self.output_dir, output_filename)
        doc.save(output_path)
        
        return output_filename
    
    def _preparar_contexto_contrato(self, locacao) -> Dict[str, Any]:
        """Preparar dados para o contrato."""
        hoje = datetime.now()
        
        return {
            # Data atual
            'data_hoje': hoje.strftime('%d/%m/%Y'),
            
            # Dados do contrato
            'numero_contrato': locacao.numero_contrato,
            'data_inicio': locacao.data_inicio.strftime('%d/%m/%Y'),
            'data_fim': locacao.data_fim.strftime('%d/%m/%Y'),
            'duracao_meses': (locacao.data_fim.year - locacao.data_inicio.year) * 12 + 
                           (locacao.data_fim.month - locacao.data_inicio.month),
            
            # Locador
            'locador_nome': locacao.imovel.locador.nome_razao_social,
            'locador_cpf_cnpj': locacao.imovel.locador.cpf_cnpj,
            'locador_endereco': getattr(locacao.imovel.locador, 'endereco', 'Não informado') or 'Não informado',
            'locador_telefone': getattr(locacao.imovel.locador, 'telefone', 'Não informado') or 'Não informado',
            'locador_email': getattr(locacao.imovel.locador, 'email', 'Não informado') or 'Não informado',
            
            # Locatário
            'locatario_nome': locacao.locatario.nome_razao_social,
            'locatario_cpf_cnpj': locacao.locatario.cpf_cnpj,
            'locatario_endereco': getattr(locacao.locatario, 'endereco', 'Não informado') or 'Não informado',
            'locatario_telefone': getattr(locacao.locatario, 'telefone', 'Não informado') or 'Não informado',
            'locatario_email': getattr(locacao.locatario, 'email', 'Não informado') or 'Não informado',
            
            # Imóvel
            'imovel_codigo': locacao.imovel.codigo_imovel,
            'imovel_tipo': locacao.imovel.get_tipo_imovel_display(),
            'imovel_endereco_completo': locacao.imovel.endereco_completo,
            'imovel_endereco': locacao.imovel.endereco,
            'imovel_numero': locacao.imovel.numero,
            'imovel_bairro': locacao.imovel.bairro,
            'imovel_cidade': locacao.imovel.cidade,
            'imovel_estado': locacao.imovel.estado,
            'imovel_cep': locacao.imovel.cep,
            
            # Valores
            'valor_aluguel': formatar_moeda_brasileira(getattr(locacao, 'valor_aluguel', None) or locacao.imovel.valor_aluguel),
            'valor_condominio': formatar_moeda_brasileira(getattr(locacao, 'valor_condominio', None) or locacao.imovel.valor_condominio or 0),
            'valor_iptu': formatar_moeda_brasileira(getattr(locacao, 'valor_iptu', None) or 0),
            'dia_vencimento': getattr(locacao, 'dia_vencimento', 5),
        }
    
    def _preparar_contexto_recibo(self, pagamento) -> Dict[str, Any]:
        """Preparar dados para o recibo."""
        return {
            'data_hoje': datetime.now().strftime('%d/%m/%Y'),
            'numero_recibo': pagamento.numero_pagamento,
            'data_pagamento': pagamento.data_pagamento.strftime('%d/%m/%Y'),
            'valor_pago': formatar_moeda_brasileira(pagamento.valor_pago),
            'forma_pagamento': pagamento.get_forma_pagamento_display(),
            
            # Comanda
            'numero_comanda': pagamento.comanda.numero_comanda,
            'mes_referencia': f"{pagamento.comanda.mes_referencia:02d}/{pagamento.comanda.ano_referencia}",
            
            # Locatário
            'locatario_nome': pagamento.comanda.locacao.locatario.nome_razao_social,
            'locatario_cpf_cnpj': pagamento.comanda.locacao.locatario.cpf_cnpj,
            
            # Imóvel
            'imovel_endereco': pagamento.comanda.locacao.imovel.endereco_completo,
            
            # Usuário
            'usuario_nome': pagamento.usuario_registro.get_full_name(),
        }
    
    def _criar_template_contrato(self, path):
        """Criar template básico de contrato."""
        doc = Document()
        
        # Título
        title = doc.add_heading('CONTRATO DE LOCAÇÃO RESIDENCIAL', 0)
        title.alignment = 1
        
        doc.add_paragraph(f'Contrato nº: {{{{ numero_contrato }}}}')
        doc.add_paragraph(f'Data: {{{{ data_hoje }}}}')
        doc.add_paragraph()
        
        # Partes
        doc.add_heading('LOCADOR', 1)
        doc.add_paragraph('Nome/Razão Social: {{ locador_nome }}')
        doc.add_paragraph('CPF/CNPJ: {{ locador_cpf_cnpj }}')
        doc.add_paragraph('Endereço: {{ locador_endereco }}')
        doc.add_paragraph('Telefone: {{ locador_telefone }}')
        doc.add_paragraph()
        
        doc.add_heading('LOCATÁRIO', 1)
        doc.add_paragraph('Nome/Razão Social: {{ locatario_nome }}')
        doc.add_paragraph('CPF/CNPJ: {{ locatario_cpf_cnpj }}')
        doc.add_paragraph('Endereço: {{ locatario_endereco }}')
        doc.add_paragraph('Telefone: {{ locatario_telefone }}')
        doc.add_paragraph()
        
        doc.add_heading('IMÓVEL LOCADO', 1)
        doc.add_paragraph('Código: {{ imovel_codigo }}')
        doc.add_paragraph('Tipo: {{ imovel_tipo }}')
        doc.add_paragraph('Endereço: {{ imovel_endereco_completo }}')
        doc.add_paragraph()
        
        doc.add_heading('CONDIÇÕES', 1)
        doc.add_paragraph('Valor do Aluguel: {{ valor_aluguel }}')
        doc.add_paragraph('Vencimento: Dia {{ dia_vencimento }} de cada mês')
        doc.add_paragraph('Período: {{ data_inicio }} até {{ data_fim }}')
        doc.add_paragraph('Duração: {{ duracao_meses }} meses')
        
        # Assinaturas
        doc.add_paragraph('\n\n\n')
        p1 = doc.add_paragraph('_' * 50)
        p1.alignment = 1
        p2 = doc.add_paragraph('LOCADOR')
        p2.alignment = 1
        
        doc.add_paragraph('\n\n')
        p3 = doc.add_paragraph('_' * 50)
        p3.alignment = 1
        p4 = doc.add_paragraph('LOCATÁRIO')
        p4.alignment = 1
        
        doc.save(path)
    
    def _criar_template_recibo(self, path):
        """Criar template básico de recibo."""
        doc = Document()
        
        title = doc.add_heading('RECIBO DE PAGAMENTO', 0)
        title.alignment = 1
        
        doc.add_paragraph(f'Recibo nº: {{{{ numero_recibo }}}}')
        doc.add_paragraph(f'Data: {{{{ data_hoje }}}}')
        doc.add_paragraph()
        
        doc.add_paragraph('Recebi de {{ locatario_nome }}, CPF/CNPJ {{ locatario_cpf_cnpj }}, '
                         'a quantia de {{ valor_pago }} referente ao pagamento de aluguel '
                         'do imóvel localizado em {{ imovel_endereco }}, '
                         'referente ao período {{ mes_referencia }}.')
        doc.add_paragraph()
        
        doc.add_paragraph('Forma de pagamento: {{ forma_pagamento }}')
        doc.add_paragraph('Data do pagamento: {{ data_pagamento }}')
        doc.add_paragraph('Comanda: {{ numero_comanda }}')
        doc.add_paragraph()
        doc.add_paragraph()
        
        doc.add_paragraph('Responsável: {{ usuario_nome }}')
        
        doc.save(path)

